from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
         "julio", "agosto", "setiembre", "octubre", "noviembre", "diciembre"]


def _fmt_fecha(fecha_str):
    try:
        d = datetime.strptime(fecha_str, "%Y-%m-%d")
        return d.strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return fecha_str


def _fmt_fecha_larga(fecha_str):
    try:
        d = datetime.strptime(fecha_str, "%Y-%m-%d")
        return f"{d.day:02d} de {MESES[d.month]} del {d.year}"
    except (ValueError, TypeError):
        return fecha_str


def _set_col_widths(table, widths):
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)


def _build_doc(doc_title, fecha_label, sections_data,
               residente_nombre, residente_cargo,
               responsable_nombre, responsable_cargo,
               asunto, proyecto_nombre, cui, clima, observaciones,
               show_fecha_col):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Header with image
    header = section.header
    header.is_linked_to_previous = False
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = os.path.join(os.path.dirname(__file__), "ENCABEZADO.png")
    if os.path.exists(img_path):
        run = p_header.add_run()
        run.add_picture(img_path, width=Cm(15))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def add_heading(text, level=1, size=10):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        return p

    def add_para(text, bold=False, size=10, space_after=Pt(4)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        return p

    def add_field(label, value_lines):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}\t: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Calibri"
        for i, line in enumerate(value_lines):
            if i > 0:
                br = p.add_run()
                br.add_break()
            r2 = p.add_run(line)
            r2.font.size = Pt(10)
            r2.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        return p

    # ── Title (underlined) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_title)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # ── Fields ──
    if residente_nombre:
        lines = [residente_nombre]
        if residente_cargo:
            lines.append(residente_cargo)
        add_field("A", lines)

    if responsable_nombre:
        lines = [responsable_nombre]
        if responsable_cargo:
            lines.append(responsable_cargo)
        add_field("DE", lines)

    if asunto:
        add_field("ASUNTO", [asunto])

    add_field("FECHA", [f"Santo Domingo de Acobamba, {fecha_label}"])

    # ── DATOS GENERALES ──
    doc.add_paragraph()
    add_heading("DATOS GENERALES", level=2, size=10)
    add_field("NOMBRE DEL PROYECTO", [proyecto_nombre])
    add_field("CUI", [cui])
    if clima:
        add_field("CLIMA", [clima])

    if observaciones:
        add_field("OBSERVACIONES", [observaciones])

    # ── Sections (partidas, materiales, notas, fotos) ──
    for section_data in sections_data:
        s_type = section_data.get("type")

        if s_type == "partidas":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            avances = section_data["data"]
            if avances:
                hdrs = ["Sector", "Código", "Partida", "Ejecutado por", "Operarios", "Oficiales", "Peones", "Horas", "Cant. Ejec."]
                if show_fecha_col:
                    hdrs = ["Fecha"] + hdrs
                table = doc.add_table(rows=1, cols=len(hdrs))
                table.style = "Light Shading Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(hdrs):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(8)
                            r.font.name = "Calibri"
                            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                if show_fecha_col:
                    col_widths = [1.5, 1.8, 1.2, 3.8, 1.8, 1.0, 1.0, 1.0, 1.0, 1.0]
                else:
                    col_widths = [1.8, 1.2, 4.5, 1.8, 1.1, 1.1, 1.1, 1.1, 1.1]

                for a in avances:
                    row = table.add_row()
                    vals = []
                    if show_fecha_col:
                        vals.append(a.get("fecha", ""))
                    vals += [
                        a.get("sector_nombre") or "",
                        a.get("partida_codigo") or "",
                        a["partida_nombre"],
                        a.get("subcontrata_nombre") or "Empresa",
                        str(a.get("num_operarios", 0)),
                        str(a.get("num_oficiales", 0)),
                        str(a.get("num_peones", 0)),
                        f'{a.get("horas_trabajadas", 0):.1f}',
                        f'{a.get("cantidad_ejecutada", 0):.2f}',
                    ]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        for p in row.cells[i].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8)
                                r.font.name = "Calibri"
                                r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                _set_col_widths(table, col_widths)
                doc.add_paragraph()

        elif s_type == "materiales":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            materiales = section_data["data"]
            if materiales:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Light Shading Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdrs = ["Material", "Cantidad", "Unidad", "Fecha"]
                for i, h in enumerate(hdrs):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(8)
                            r.font.name = "Calibri"
                            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                for mat in materiales:
                    row = table.add_row()
                    vals = [mat["material"], str(mat["cantidad"]), mat["unidad"], mat["fecha"]]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        for p in row.cells[i].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8)
                                r.font.name = "Calibri"
                                r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                doc.add_paragraph()

        elif s_type == "notas":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            for nota in section_data["data"]:
                add_para(f"- {nota['nota']}", size=10)

        elif s_type == "fotos":
            doc.add_paragraph()
            add_heading(section_data["title"], level=2, size=10)
            for idx, foto in enumerate(section_data["data"], start=1):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"FOTOGRAFIA N° {idx}")
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Calibri"
                if os.path.exists(foto["ruta"]):
                    try:
                        doc.add_picture(foto["ruta"], width=Inches(4.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        p2 = doc.add_paragraph()
                        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r2 = p2.add_run("[No se pudo insertar la imagen]")
                        r2.font.size = Pt(9)
                if foto["descripcion"]:
                    p3 = doc.add_paragraph()
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r3 = p3.add_run(f"Descripcion: {foto['descripcion']}")
                    r3.font.size = Pt(10)
                    r3.font.name = "Calibri"

    # ── Signatures ──
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Signature ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_________________________________")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(responsable_nombre)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(responsable_cargo)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    return doc


def generar_informe(fecha, avances, materiales, notas, fotos, registro,
                    responsable_nombre, responsable_cargo,
                    residente_nombre=None, residente_cargo=None,
                    asunto="", proyecto_nombre="", cui=""):
    fecha_label = _fmt_fecha_larga(fecha)
    sections = []
    if avances:
        sections.append({"type": "partidas", "title": "AVANCE DE PARTIDAS", "data": avances})
    if materiales:
        sections.append({"type": "materiales", "title": "CONTROL DE MATERIALES", "data": materiales})
    if notas:
        sections.append({"type": "notas", "title": "NOTAS DE CAMPO", "data": notas})
    if fotos:
        sections.append({"type": "fotos", "title": "REGISTRO FOTOGRAFICO", "data": fotos})

    clima = registro.get("clima", "") if registro else ""
    observaciones = registro.get("observaciones", "") if registro else ""

    return _build_doc("REPORTE DIARIO DE OBRA", fecha_label, sections,
                      residente_nombre, residente_cargo,
                      responsable_nombre, responsable_cargo,
                      asunto, proyecto_nombre, cui, clima, observaciones,
                      show_fecha_col=False)


def generar_informe_semanal(fecha_ini, fecha_fin, registros, avances, materiales, notas,
                            responsable_nombre, responsable_cargo,
                            residente_nombre=None, residente_cargo=None,
                            asunto="", proyecto_nombre="", cui=""):
    fecha_label = f"Del {_fmt_fecha(fecha_ini)} al {_fmt_fecha(fecha_fin)}"
    sections = []

    clima = "Varios (ver detalle)"
    observaciones = ""

    if registros:
        lineas = []
        for r in registros:
            lineas.append(f"Dia {_fmt_fecha(r['fecha'])}: Clima {r['clima']}")
        observaciones = f"Dias registrados: {len(registros)}\n" + "\n".join(lineas)

    if avances:
        sections.append({"type": "partidas", "title": "AVANCE DE PARTIDAS", "data": avances})
    if materiales:
        sections.append({"type": "materiales", "title": "CONTROL DE MATERIALES", "data": materiales})
    if notas:
        sections.append({"type": "notas", "title": "NOTAS DE CAMPO", "data": notas})

    return _build_doc("REPORTE SEMANAL DE OBRA", fecha_label, sections,
                      residente_nombre, residente_cargo,
                      responsable_nombre, responsable_cargo,
                      asunto, proyecto_nombre, cui, clima, observaciones,
                      show_fecha_col=True)
